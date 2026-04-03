from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)

def add_text(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_formula(label, formula, explanation):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(label)
    run.bold = True
    run = p.add_run(f"  {formula}")
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.5)
    run2 = p2.add_run(explanation)
    run2.italic = True
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_matrix(label, rows, col_headers=None, row_headers=None, note=None):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True

    num_cols = len(rows[0]) + (1 if row_headers else 0)
    num_rows = len(rows) + (1 if col_headers else 0)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    start_row = 0
    start_col = 0

    if col_headers:
        offset = 1 if row_headers else 0
        for j, h in enumerate(col_headers):
            cell = table.cell(0, j + offset)
            cell.text = h
            for par in cell.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.bold = True
        start_row = 1

    if row_headers:
        for i, h in enumerate(row_headers):
            cell = table.cell(i + start_row, 0)
            cell.text = h
            for par in cell.paragraphs:
                for run in par.runs:
                    run.bold = True
        start_col = 1

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i + start_row, j + start_col)
            cell.text = str(val)
            for par in cell.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.font.name = "Consolas"

    if note:
        p2 = doc.add_paragraph()
        run2 = p2.add_run(note)
        run2.italic = True
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

add_heading("GP Regression — Fully Worked Example", level=0)

add_heading("Setup", level=1)
add_text("One soil sample. We measured reflectance at 2 wavelengths. We want to predict reflectance at 3 new wavelengths.")

add_matrix(
    "Training Data (what we observed):",
    [["400 nm", "0.10"], ["600 nm", "0.30"]],
    col_headers=["Wavelength", "Reflectance"]
)

add_matrix(
    "Prediction Points (what we want to know):",
    [["450 nm", "?"], ["500 nm", "?"], ["550 nm", "?"]],
    col_headers=["Wavelength", "Reflectance"]
)

add_text("Kernel: RBF with length_scale (ℓ) = 100 nm", bold=True)
add_formula("Formula:", "k(x, x') = exp( −0.5 × ((x − x') / ℓ)² )",
            "\"How similar are two wavelengths? Closer → more similar → higher value.\"")

add_heading("Step 1: K_train_train — How Observed Points Relate to Each Other", level=1)
add_text("Compute kernel between every pair of observed wavelengths. This gives a 2×2 matrix.")

add_formula("k(400, 400):", "exp(−0.5 × (0/100)²)   = exp(0)     = 1.0000", "Same point → perfect correlation")
add_formula("k(400, 600):", "exp(−0.5 × (200/100)²) = exp(−2.0) = 0.1353", "200nm apart → weak correlation")
add_formula("k(600, 400):", "= 0.1353", "Symmetric — same as above")
add_formula("k(600, 600):", "exp(−0.5 × (0/100)²)   = exp(0)     = 1.0000", "Same point → perfect correlation")

add_matrix(
    "K_train_train (2×2):",
    [["1.0000", "0.1353"], ["0.1353", "1.0000"]],
    col_headers=["400nm", "600nm"],
    row_headers=["400nm", "600nm"],
    note="Reading: 400nm and 600nm are 200nm apart, so correlation is only 0.14 — they don't know much about each other."
)

add_heading("Step 2: K_pred_train — How Prediction Points Relate to Observed Points", level=1)
add_text("This is the most important matrix. It tells us how much each prediction point 'listens to' each observation.")

add_formula("k(450, 400):", "exp(−0.5 × (50/100)²)  = exp(−0.125) = 0.8825", "450 is close to 400 → high similarity")
add_formula("k(450, 600):", "exp(−0.5 × (150/100)²) = exp(−1.125) = 0.3247", "450 is far from 600 → low similarity")
doc.add_paragraph()
add_formula("k(500, 400):", "exp(−0.5 × (100/100)²) = exp(−0.500) = 0.6065", "500 is medium distance from 400")
add_formula("k(500, 600):", "exp(−0.5 × (100/100)²) = exp(−0.500) = 0.6065", "500 is equally far from 600")
doc.add_paragraph()
add_formula("k(550, 400):", "exp(−0.5 × (150/100)²) = exp(−1.125) = 0.3247", "550 is far from 400 → low similarity")
add_formula("k(550, 600):", "exp(−0.5 × (50/100)²)  = exp(−0.125) = 0.8825", "550 is close to 600 → high similarity")

add_matrix(
    "K_pred_train (3×2):",
    [["0.8825", "0.3247"], ["0.6065", "0.6065"], ["0.3247", "0.8825"]],
    col_headers=["400nm", "600nm"],
    row_headers=["450nm", "500nm", "550nm"],
    note="Read each row: 450nm mostly listens to 400nm. 500nm listens equally. 550nm mostly listens to 600nm."
)

add_heading("Step 3: K_pred_pred — How Prediction Points Relate to Each Other", level=1)
add_text("Needed for computing uncertainty. Same kernel, applied between prediction points.")

add_matrix(
    "K_pred_pred (3×3):",
    [["1.0000", "0.8825", "0.6065"],
     ["0.8825", "1.0000", "0.8825"],
     ["0.6065", "0.8825", "1.0000"]],
    col_headers=["450nm", "500nm", "550nm"],
    row_headers=["450nm", "500nm", "550nm"],
    note="This represents our prior uncertainty — before seeing any data."
)

add_heading("Step 4: Invert K_train_train", level=1)
add_text("Standard 2×2 matrix inversion.")
add_formula("det:", "1.0 × 1.0 − 0.1353 × 0.1353 = 0.9817", "")

add_matrix(
    "K_inv (2×2):",
    [["1.0186", "−0.1378"], ["−0.1378", "1.0186"]],
    col_headers=["", ""],
)

add_heading("Step 5: Compute Posterior Mean — THE PREDICTION", level=1)
add_text("Formula:  posterior_mean = K_pred_train  ×  K_inv  ×  y_train", bold=True)

add_text("Sub-step A:  alpha = K_inv × y_train", bold=True)
add_formula("alpha[0]:", "1.0186 × 0.10 + (−0.1378) × 0.30 = 0.0605", "\"Weight\" from the 400nm observation")
add_formula("alpha[1]:", "(−0.1378) × 0.10 + 1.0186 × 0.30 = 0.2918", "\"Weight\" from the 600nm observation")

add_text("Sub-step B:  posterior_mean = K_pred_train × alpha", bold=True)
add_formula("At 450nm:", "0.8825 × 0.0605  +  0.3247 × 0.2918  =  0.1482",
            "Close to 400nm (0.10) → pulled low")
add_formula("At 500nm:", "0.6065 × 0.0605  +  0.6065 × 0.2918  =  0.2137",
            "Equidistant → right between the two values")
add_formula("At 550nm:", "0.3247 × 0.0605  +  0.8825 × 0.2918  =  0.2772",
            "Close to 600nm (0.30) → pulled high")

add_matrix(
    "Predictions:",
    [["450 nm", "0.1482"], ["500 nm", "0.2137"], ["550 nm", "0.2772"]],
    col_headers=["Wavelength", "Predicted Reflectance"],
    note="The GP smoothly interpolated between 0.10 (at 400nm) and 0.30 (at 600nm), weighting by proximity."
)

add_heading("Step 6: Compute Posterior Covariance — THE UNCERTAINTY", level=1)
add_text("Formula:  posterior_cov = K_pred_pred − K_pred_train × K_inv × K_pred_train^T", bold=True)
add_text("Meaning: start with prior uncertainty, subtract what we learned from data.")

add_matrix(
    "Posterior Covariance (3×3):",
    [["0.1783", "0.2374", "0.1445"],
     ["0.2376", "0.3520", "0.2376"],
     ["0.1445", "0.2374", "0.1783"]],
    col_headers=["450nm", "500nm", "550nm"],
    row_headers=["450nm", "500nm", "550nm"],
    note="The diagonal gives variance at each point. Off-diagonals tell how uncertainties are correlated."
)

add_text("Standard deviations (square root of diagonal):", bold=True)
add_matrix(
    "",
    [["450 nm", "0.1783", "0.4222"],
     ["500 nm", "0.3520", "0.5933"],
     ["550 nm", "0.1783", "0.4222"]],
    col_headers=["Wavelength", "Variance", "Std Dev (σ)"],
    note="500nm has the MOST uncertainty (σ=0.59) because it's farthest from both observations. "
         "450nm and 550nm have LESS uncertainty (σ=0.42) because each is close to an observed point."
)

add_heading("Final Summary", level=1)

add_matrix(
    "",
    [["400 nm", "0.10", "— (observed)"],
     ["450 nm", "0.1482", "± 0.84"],
     ["500 nm", "0.2137", "± 1.19"],
     ["550 nm", "0.2772", "± 0.84"],
     ["600 nm", "0.30", "— (observed)"]],
    col_headers=["Wavelength", "Reflectance", "95% CI (±2σ)"],
)

p = doc.add_paragraph()
run = p.add_run("The GP smoothly interpolated between observations, with uncertainty that is "
                "smallest near observed data and largest in the gap between them.")
run.italic = True

output_path = "gp_regression_worked_example.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
